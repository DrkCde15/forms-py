from flask import Flask, request, render_template, send_file, jsonify, session
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import google.generativeai as genai
from dotenv import load_dotenv
import os
import tempfile
import secrets
import logging
from datetime import datetime

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

# Configurar API
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("Configure a variável de ambiente GEMINI_API_KEY")
genai.configure(api_key=api_key)

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", secrets.token_hex(16))

# Caminho do currículo base
BASE_CURRICULO = r"C:\Users\Júlio César\Documents\Currículo\Currículo Stephanie Amarante - Área TI (1).docx"

def verificar_arquivo_base():
    """Verifica se o arquivo base existe"""
    if not os.path.exists(BASE_CURRICULO):
        logger.error(f"Arquivo base não encontrado: {BASE_CURRICULO}")
        return False
    return True

def carregar_curriculo_base():
    """Lê o currículo base e retorna o texto"""
    try:
        if not verificar_arquivo_base():
            raise FileNotFoundError("Arquivo base do currículo não encontrado")
        
        doc = Document(BASE_CURRICULO)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        if not texto.strip():
            raise ValueError("Arquivo base está vazio ou não pôde ser lido")
            
        return texto
    except Exception as e:
        logger.error(f"Erro ao carregar currículo base: {e}")
        raise

def validar_dados_entrada(nome, resumo, objetivo):
    """Valida os dados de entrada"""
    erros = []
    
    if not nome or len(nome.strip()) < 2:
        erros.append("Nome deve ter pelo menos 2 caracteres")
    
    if not resumo or len(resumo.strip()) < 10:
        erros.append("Resumo profissional deve ter pelo menos 10 caracteres")
    
    if not objetivo or len(objetivo.strip()) < 10:
        erros.append("Objetivo profissional deve ter pelo menos 10 caracteres")
    
    # Verificar se não são apenas espaços ou caracteres especiais
    if nome and not any(c.isalpha() for c in nome):
        erros.append("Nome deve conter pelo menos uma letra")
    
    return erros

def criar_documento_formatado(texto_gerado, nome):
    """Cria um documento Word com formatação adequada"""
    novo_doc = Document()
    
    # Configurar margens
    sections = novo_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    linhas = texto_gerado.split("\n")
    primeira_linha = True
    
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue
            
        # Nome principal (primeira linha significativa)
        if primeira_linha and nome.upper() in linha.upper():
            p = novo_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(linha)
            run.bold = True
            run.font.size = Inches(0.2)
            primeira_linha = False
            
        # Seções principais (títulos)
        elif any(x in linha.upper() for x in ["RESUMO", "OBJETIVO", "EXPERIÊNCIA", "FORMAÇÃO", "HABILIDADES", "CONTATO"]) and len(linha) < 50:
            p = novo_doc.add_paragraph()
            p.add_run().add_break()  # Espaço antes da seção
            run = p.add_run(linha)
            run.bold = True
            run.underline = True
            
        # Texto em maiúsculas (subtítulos)
        elif linha.isupper() and len(linha) < 80:
            p = novo_doc.add_paragraph()
            run = p.add_run(linha)
            run.bold = True
            
        # Texto normal
        else:
            novo_doc.add_paragraph(linha)
    
    return novo_doc

def gerar_curriculo_arquivo(nome, resumo, objetivo, experiencias="", educacao="", habilidades="", linkedin="", email="", github="", idiomas=""):
    """Gera arquivo de currículo personalizado"""
    try:
        # Validar dados obrigatórios
        erros = validar_dados_entrada(nome, resumo, objetivo)
        if erros:
            raise ValueError("; ".join(erros))

        # Base da Stephanie
        doc_texto = carregar_curriculo_base()

        # Prompt com todos os dados
        prompt = f"""
Você é um especialista em RH e criação de currículos.

Use o documento de referência como modelo de estilo:
{doc_texto}

Crie um currículo completo para o usuário com os seguintes dados:

- Nome: {nome}
- Resumo Profissional: {resumo}
- Objetivo Profissional: {objetivo}
- Experiências: {experiencias if experiencias else "A definir"}
- Formação: {educacao if educacao else "A informar"}
- Habilidades: {habilidades if habilidades else "Versatilidade e aprendizado contínuo"}
- LinkedIn: {linkedin if linkedin else "Não informado"}
- Email: {email if email else "Não informado"}
- GitHub: {github if github else "Não informado"}
- Idiomas: {idiomas if idiomas else "Não informado"}

Instruções:
1. Mantenha a estrutura clara e profissional
2. Inclua seções bem definidas
3. Use linguagem formal e objetiva
4. Currículo entre 400-800 palavras
"""

        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        texto_final = response.text

        # Criar documento formatado
        novo_doc = criar_documento_formatado(texto_final, nome)

        # Salvar arquivos
        tmp_dir = tempfile.mkdtemp()
        nome_arquivo = nome.replace(' ', '_').replace('/', '_').replace('\\', '_')
        output_docx = os.path.join(tmp_dir, f"{nome_arquivo}_Curriculo.docx")
        output_pdf = os.path.join(tmp_dir, f"{nome_arquivo}_Curriculo.pdf")

        novo_doc.save(output_docx)

        try:
            convert(output_docx, output_pdf)
        except Exception as pdf_error:
            logger.warning(f"Erro na conversão para PDF: {pdf_error}")
            output_pdf = None

        return output_docx, output_pdf

    except Exception as e:
        logger.error(f"Erro na geração do currículo: {e}")
        raise


@app.route("/")
def index():
    """Página principal"""
    return render_template("index.html")

@app.route("/gerar_curriculo", methods=["POST"])
def gerar_curriculo_route():
    try:
        idiomas_nomes = request.form.getlist("idioma_nome[]")
        idiomas_niveis = request.form.getlist("idioma_nivel[]")
        idiomas = [f"{nome} ({nivel})" for nome, nivel in zip(idiomas_nomes, idiomas_niveis) if nome.strip()]

        dados = {
            'nome': request.form.get("nome", "").strip(),
            'resumo': request.form.get("resumo", "").strip(),
            'objetivo': request.form.get("objetivo", "").strip(),
            'experiencias': request.form.get("experiencias", "").strip(),
            'educacao': request.form.get("educacao", "").strip(),
            'habilidades': request.form.get("habilidades", "").strip(),
            'linkedin': request.form.get("linkedin", "").strip(),
            'email': request.form.get("email", "").strip(),
            'github': request.form.get("github", "").strip(),
            'idiomas': ", ".join(idiomas),
        }

        docx_path, pdf_path = gerar_curriculo_arquivo(**dados)

        if pdf_path and os.path.exists(pdf_path):
            return send_file(
                pdf_path,
                as_attachment=True,
                download_name=f"{dados['nome'].replace(' ', '_')}_Curriculo.pdf"
            )
        else:
            return send_file(
                docx_path,
                as_attachment=True,
                download_name=f"{dados['nome'].replace(' ', '_')}_Curriculo.docx"
            )

    except ValueError as ve:
        logger.warning(f"Erro de validação: {ve}")
        return jsonify({'error': str(ve)}), 400
    except Exception as e:
        logger.error(f"Erro interno: {e}")
        return jsonify({'error': 'Erro interno do servidor. Tente novamente.'}), 500


@app.route("/download_pdf")
def download_pdf():
    """Endpoint para download do PDF"""
    try:
        ultimo_curriculo = session.get('ultimo_curriculo')
        if not ultimo_curriculo or not ultimo_curriculo.get('pdf'):
            return jsonify({'error': 'PDF não disponível'}), 404
        
        pdf_path = ultimo_curriculo['pdf']
        if not os.path.exists(pdf_path):
            return jsonify({'error': 'Arquivo PDF não encontrado'}), 404
        
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"{ultimo_curriculo['nome'].replace(' ', '_')}_Curriculo.pdf"
        )
        
    except Exception as e:
        logger.error(f"Erro no download PDF: {e}")
        return jsonify({'error': 'Erro ao baixar PDF'}), 500

@app.route("/health")
def health_check():
    """Endpoint de verificação de saúde"""
    try:
        # Verificar se o arquivo base existe
        if not verificar_arquivo_base():
            return jsonify({'status': 'error', 'message': 'Arquivo base não encontrado'}), 500
        
        # Verificar API do Gemini
        model = genai.GenerativeModel("gemini-1.5-flash")
        test_response = model.generate_content("Test")
        
        return jsonify({
            'status': 'ok',
            'timestamp': datetime.now().isoformat(),
            'gemini_api': 'connected'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint não encontrado'}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Erro interno: {error}")
    return jsonify({'error': 'Erro interno do servidor'}), 500

if __name__ == "__main__":
    # Verificações iniciais
    if not verificar_arquivo_base():
        print("AVISO: Arquivo base do currículo não encontrado!")
        print(f"Caminho esperado: {BASE_CURRICULO}")
    
    app.run(debug=True, host='127.0.0.1', port=5000)